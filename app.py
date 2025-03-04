# Define the algorithm as a string with Markdown formatting
import os
import streamlit as st

st.set_page_config(
    layout="wide"
)

algorithm_code = """
@tool("EDA Tool")
def eda_tool(data: str) -> str:
    return "Performed EDA on {data}"

@tool("Code Executor")
def code_executor(code: str) -> str:
    result = execute(code)
    return result

manager = Agent(
    role="Data Science Manager",
    goal="Manage the overall modeling pipeline with specific instructions for each agent.",
    expertise="An expert skilled in managing a modeling team.",
    tools=[eda_tool, code_executor],
    memory=True
)

eda_agent = Agent(
    role="Senior Data Scientist",
    goal="Conduct a detailed exploratory data analysis.",
    expertise="An EDA expert with experience in data visualization and insights extraction.",
    tools=[eda_tool]
)

feature_engineer = Agent(
    role="Senior Data Scientist",
    goal="Preprocess the data for model training.",
    expertise="A skilled engineer with expertise in data transformation and selection.",
    tools=[code_executor]
)

model_selection_agent = Agent(
    role="Machine Learning Engineer",
    goal="Select the best Machine Learning model.",
    expertise="An expert in using GridSearch and RandomSearch for model selection.",
    tools=[code_executor]
)

hyperparameter_tuning_agent = Agent(
    role="Senior Machine Learning Engineer",
    goal="Get the optimal hyperparameters for the selected model.",
    expertise="A hyperparameter tuning expert specializing in optimization techniques.",
    tools=[code_executor]
)

model_training_agent = Agent(
    role="Senior Machine Learning Engineer",
    goal="Train the selected model with the best performing hyperparameters.",
    expertise="A machine learning training expert who fine-tunes models for high accuracy.",
    tools=[code_executor]
)

model_evaluation_agent = Agent(
    role="Senior Machine Learning Engineer",
    goal="Evaluate the trained model.",
    expertise="A model evaluation expert skilled in assessing performance metrics.",
    tools=[code_executor]
)

"""

# Set the title of the app
st.title("Agentic System Demo - Research Team")

# Create tabs
tab1, tab2 = st.tabs(["Algorithm", "System Demo"])

# Content for Tab 1
with tab1:
    st.header("Agentic System Algorithm")
    # st.write("This is the first tab. You can display any content here, such as text, images, or graphs.")
    st.code(algorithm_code, language='python')  # Display as code

# Content for Tab 2
with tab2:
    # st.header("Content for Tab 2")
    # st.write("This is the second tab. You can also display different types of content here.")
    # st.line_chart([1, 2, 3, 4, 5], use_container_width=True)
    from crewai import Agent, Task, Crew, Process
    import os
    from crewai_tools import tool
    from crewai_tools import ScrapeWebsiteTool, SerperDevTool
    from dotenv import load_dotenv
    from langchain_openai import ChatOpenAI
    from docx import Document
    from io import BytesIO
    import base64

    import logging
    from langchain_core.callbacks.base import BaseCallbackHandler
    from langchain_core.callbacks.manager import CallbackManager
    import time

    ##############################################################################
    # Suppress warning from langchain callback
    logging.getLogger("langchain_core.callbacks.manager").setLevel(logging.ERROR)

    # Warning control -- Suppress some warnings
    import warnings

    #Ignore generic warnings
    warnings.filterwarnings('ignore')

    # Suppress convergence warning from LogisticRegression
    from sklearn.exceptions import ConvergenceWarning
    warnings.simplefilter("ignore", ConvergenceWarning)
    ##############################################################################

    ##############################################################################
    # Load Environment Variables
    load_dotenv()
    ##############################################################################


    ##############################################################################
    # LLM object and API Key
    os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
    os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")
    ##############################################################################

    ##############################################################################
    # Initialize some tools
    search_tool = SerperDevTool()
    scrape_tool = ScrapeWebsiteTool()

    # Choose the llm to use
    llm = ChatOpenAI(
        model="gpt-3.5-turbo",
        temperature=0.3,
    )
    ##############################################################################

    ##############################################################################
    # EDA Tool
    # Import necessary libraries
    import pandas as pd
    import numpy as np
    import matplotlib.pyplot as plt
    import seaborn as sns
    from collections import Counter

    # The function
    def exploratory_data_analysis(df):
        report = []

        # Data overview
        total_rows, total_columns = df.shape
        report.append(f"The data contains {total_rows} rows and {total_columns} columns.")
        report.append(f"The target variable is the {df.columns[-1]} column.")

        # Check for missing values
        missing_values = df.isnull().sum()
        missing_features = missing_values[missing_values > 0]
        missing_count = len(missing_features)
        report.append(f"Out of all the features, {missing_count} features have missing values.")

        # Check for class imbalance
        class_proportions = df[df.columns[-1]].value_counts(normalize=True)
        most_common = Counter(class_proportions).most_common()[0]
        least_common = Counter(class_proportions).most_common()[-1]
        if most_common[0] > 0.7:
          report.append(f"The dataset is highly imbalanced because, the {df.columns[-1]} feature with value {str(most_common[1])} has a proportion of {str(round(most_common[0]*100, 2))}%.")
        elif most_common[0] > 0.6:
          report.append(f"The dataset is slightly imbalanced because, the {df.columns[-1]} feature with value {str(most_common[1])} has a proportion of {str(round(most_common[0]*100, 2))}%.")
        else:
          report.append(f"The dataset is balanced")

        for feature in missing_features.index:
            missing_percentage = (missing_values[feature] / total_rows) * 100
            report.append(f"The feature '{feature}' has {missing_values[feature]} missing values, which is {missing_percentage:.2f}% of the total.")

        # Identify feature types
        numeric_features = df.select_dtypes(include=[np.number]).columns.tolist()
        numeric_features = [item for item in numeric_features if item != df.columns[-1]]
        date_features = df.select_dtypes(include=[np.datetime64]).columns.tolist()
        date_features = [item for item in date_features if item != df.columns[-1]]
        text_features = df.select_dtypes(include=[object]).columns.tolist()
        text_features = [item for item in text_features if item != df.columns[-1]]

        numeric_cat = (df[numeric_features]
                        .apply(pd.Series.nunique)
                        .loc[lambda x: x <= 20])
        numeric_cat = numeric_cat.index.tolist()
        object_cat = (df[text_features]
                        .apply(pd.Series.nunique)
                        .loc[lambda x: x <= 20])
        object_cat = object_cat.index.tolist()

        all_cat = numeric_cat + object_cat
        text_features = list(set(text_features) - set(object_cat))

        report.append(f"The numeric features are: {', '.join(numeric_features)}.")
        if date_features:
            report.append(f"The date features are: {', '.join(date_features)}.")
        else:
            report.append(f"There are no date features.")
        if all_cat:
            report.append(f"The categorical features are: {', '.join(all_cat)}.")
        else:
            report.append(f"There are no categorical features.")
        if object_cat:
            report.append(f"The categorical features that needs label encoding are: {', '.join(object_cat)}.")
        else:
            report.append(f"There are no categorical features that needs label encoding.")
        if text_features:
            report.append(f"The text features are: {', '.join(text_features)}.")
        else:
            report.append(f"There are no text features that are standalone.")

        # Descriptive statistics for numeric features
        descriptive_stats_dict = {}

        for feature in numeric_features:
            desc = df[feature].describe()
            descriptive_stats_dict[feature] = desc.to_dict()

            # Checking skewness
            skewness = df[feature].skew()
            if skewness > 0:
                report.append(f"The feature '{feature}' is positively skewed.")
            elif skewness < 0:
                report.append(f"The feature '{feature}' is negatively skewed.")
            else:
                report.append(f"The feature '{feature}' has a normal distribution.")

        report.append("Descriptive statistics for numeric features:")
        report.append(str(descriptive_stats_dict))
        # Correlation matrix for numeric features
        correlation_matrix = df[numeric_features].corr()
        highly_correlated = correlation_matrix[(correlation_matrix.abs() > 0.8) & (correlation_matrix != 1.0)]

        # if not highly_correlated.empty:
        if not highly_correlated.dropna(axis=1, how='all').empty: # If their is high correlation features
            for feature in highly_correlated.columns:
                correlated_features = highly_correlated[feature][highly_correlated[feature].abs() > 0.5].index.tolist()
                if correlated_features:
                  report.append(f"The feature '{feature}' is highly correlated with: {', '.join(correlated_features)}.")
        else:
            report.append("There are no highly correlated features.")

        return "\n".join(report)


    # Making the EDA function a crewai tool
    @tool("Exploratory Data Analysis Tool")
    def eda_tool(csv_path: str) -> str:
        """
        Performs exploratory data analysis on a given CSV file.

        Args:
            csv_path (str): The path to the CSV file.

        Returns:
            str: A report summarizing the exploratory data analysis.
        """
        # Load the CSV file into a Pandas DataFrame
        df = pd.read_csv(csv_path)

        # Call the exploratory_data_analysis function
        report = exploratory_data_analysis(df)

        return report
    ##############################################################################

    ##############################################################################
    # Code Execution tool
    # Import required libraries
    import re

    # Function to escape newlines in double-quoted strings
    def escape_newlines_in_quotes(code: str) -> str:
        """Escape newline characters in double-quoted strings."""
        # Use regex to find double-quoted strings and replace \n within those strings
        def replace_newlines(match):
            return match.group(0).replace("\n", "\\n")

        # This pattern matches double-quoted strings, capturing them
        pattern = r'"([^"\\]*(?:\\.[^"\\]*)*)"'  # Matches double-quoted strings
        return re.sub(pattern, replace_newlines, code)

    # Function to execute the generated code
    def execute_generated_code(generated_code):
      generated_code = escape_newlines_in_quotes(generated_code)
      generated_code = generated_code.replace("\\", "\\\\")
      local_vars = {}
      # exec(generated_code, {}, local_vars)
      exec(generated_code, globals(), local_vars)
      return local_vars.get('result', None)

    # Initializing the function as a crewai tool
    @tool("Python Code Executor Tool")
    def code_execution_tool(generated_code: str) -> str:
        """
        Executes the code and returns the output as str.

        Args:
            generated_code (str): The code that needs to be executed.

        Returns:
            str: The output of the executed code.
        """
        # Execute the generated code
        output = execute_generated_code(generated_code)

        return output
    ##############################################################################

    ##############################################################################
    # RAG tool
    # import neccessary libraries
    from pypdf import PdfReader

    def rag_format(pdf_path, model_doc_path):
        pdf_content = ""

        # PDF reader object
        reader = PdfReader(pdf_path)

        # Get the content of all the pages
        for i in range(len(reader.pages)):
          pdf_content += reader.pages[i].extract_text() + "\n\n"

        pdf_content = pdf_content.replace("\t", " ")

        with open(model_doc_path, 'r') as file:
          all_model_output = json.load(file)

        # Get the modeling documentation
        model_output = all_model_output[0]['raw']  + "\n\n" +  all_model_output[1]['raw'] + "\n\n" +  all_model_output[2]['raw'] + "\n\n" +  all_model_output[3]['raw'] + "\n\n" +  all_model_output[4]['raw'] + "\n\n" +  all_model_output[5]['raw']

        # Prepare a prompt for the LLM
        prompt = (
            f"##Reading modeling documentation from the modeling team....\n"
            f"*Modeling Documentation*\n"
            f"{model_output}\n\n"
            f"End of modeling documentation.....\n\n\n"
            f"##Reading the organizational modeling guide....\n"
            f"*Comprehensive Guide to Machine Learning Pipeline Development*\n\n"
            f"{pdf_content}\n\n"
            f"End of the organizational modeling guide or blueprint.....\n\n\n"
            f"Determine if the modeling documentation adhered to the processes described in the organizational modeling guide. "
        )

        return prompt


    # Creating RAG tool with crewai tool instance
    @tool("Retrieval Augmented Generation(RAG) Tool")
    def rag_tool(pdf_path: str, model_doc_path: str) -> str:
        """
        Performs RAG on two documents to verify authenticity.

        Args:
            pdf_path (str): The path to the organizational modeling guide PDF file.
            model_doc_path (str): The path to the modeling documentation text file.

        Returns:
            str: A formatted form of the two documents for further processing.
        """

        # Call the rag_format function
        report_format = rag_format(pdf_path, model_doc_path)

        return report_format
    ##############################################################################

    ##############################################################################
    # Model selection tool
    import pandas as pd
    from sklearn.model_selection import train_test_split
    from sklearn.model_selection import GridSearchCV
    from sklearn.linear_model import LogisticRegression
    from sklearn.tree import DecisionTreeClassifier
    from sklearn.svm import SVC
    from sklearn.naive_bayes import GaussianNB
    from sklearn.ensemble import RandomForestClassifier
    from xgboost import XGBClassifier
    from sklearn.ensemble import GradientBoostingClassifier
    from catboost import CatBoostClassifier
    from sklearn.ensemble import AdaBoostClassifier

    def select_best_model(csv_path, target):

      # Load the dataset
      if csv_path.startswith("fraud_und"): # For fraud undersampled
        data = pd.read_csv(csv_path)
      elif csv_path.startswith("c"): # For credit and card approval
        data = pd.read_csv(csv_path)

        # Generate a random seed
        np.random.seed(42)

        # Calculate the number of rows to select
        n_rows = int(0.05 * len(data)) # 5% of the dataset

        # Generate a random subset of indices
        indices = np.random.choice(len(data), n_rows, replace=False)

        # Select the subset of the dataframe
        data = data.iloc[indices]
      else:
        data = pd.read_csv(csv_path)

        # Generate a random seed
        np.random.seed(42)

        # Calculate the number of rows to select
        n_rows = int(0.005 * len(data)) # 0.5% of the dataset

        # Generate a random subset of indices
        indices = np.random.choice(len(data), n_rows, replace=False)

        # Select the subset of the dataframe
        data = data.iloc[indices]

      X = data.drop(target, axis=1)
      y = data[target]

      # Split the data into training and testing sets
      X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

      # Define the models
      models = [
          ('Logistic Regression', LogisticRegression(random_state=42, max_iter=1000)),
          ('Decision Tree', DecisionTreeClassifier(random_state=42)),
          ('Support Vector Machine', SVC(random_state=42)),
          ('Naive Bayes', GaussianNB()),
          ('Random Forest', RandomForestClassifier(random_state=42)),
          ('XGBoost', XGBClassifier(random_state=42)),
          ('Gradient Boosting', GradientBoostingClassifier(random_state=42)),
          ('CatBoost', CatBoostClassifier(verbose=0,random_state=42)),
          ('AdaBoosting', AdaBoostClassifier(random_state=42))
      ]

      # GridSearchCV to select the best model
      best_model = None
      best_score = 0
      for name, model in models:
          grid_search = GridSearchCV(model, param_grid={}, cv=5, n_jobs=-1)
          grid_search.fit(X_train, y_train)
          score = grid_search.best_score_
          if score > best_score:
              best_score = score
              best_model = name

      return best_model


    # Creating crewai tool
    @tool("Model Selection Tool")
    def mod_sel_tool(csv_path: str, target: str) -> str:
        """
        Performs model selection using GridSearchCV on a given CSV file.

        Args:
            csv_path (str): The path to the CSV file.
            target (str): The name of the target variable.

        Returns:
            str: The best performing machine learning model for the provided CSV file.
        """
        # Load the CSV file into a Pandas DataFrame
        print('Reading the CSV file before performing grid search......')

        print('I have completely read the CSV file and now conducting model selection using GridSearchCV.....')
        # Call the model selection function
        model_selected = select_best_model(csv_path, target)

        return "The selected model based on performance on the train dataset is the " + model_selected.upper() + " algorithm."
    ##############################################################################

    ##############################################################################
    # Store logs in a list
    log_messages = []


    # Helper function to log messages to file and console
    def log(msg):
        # print(msg)  # Print to console
        st.write(msg) # Print to streamlit
        log_messages.append(msg)  # Store in a list
    ##############################################################################


    ##############################################################################
    # Custom callback handler to capture CrewAI execution
    class CrewAICallbackHandler(BaseCallbackHandler):
        """Callback handler to track all CrewAI actions."""

        def on_chain_start(self, serialized, inputs, **kwargs):
            log(f"🚀 Task Started: {serialized} | Inputs: {inputs}")

        def on_chain_end(self, outputs, **kwargs):
            log(f"✅ Task Completed | Outputs: {outputs}")

        def on_tool_start(self, serialized, inputs, **kwargs):
            log(f"🛠️ Tool Invoked: {serialized} | Inputs: {inputs}")

        def on_tool_end(self, outputs, **kwargs):
            log(f"✅ Tool Result: {outputs}")

        def on_agent_action(self, action, **kwargs):
            log(f"🤖 Agent Action: {action}")

    class ExtendedCallbackManager(CallbackManager):
        def on_chain_end(self, outputs, **kwargs):

            # Invoke the method from CrewAICallbackHandler
            for callback in self.handlers:
                if hasattr(callback, 'on_chain_end'):
                    callback.on_chain_end(outputs, **kwargs)

        def on_agent_action(self, action, **kwargs):

            # Invoke the method from CrewAICallbackHandler
            for callback in self.handlers:
                if hasattr(callback, 'on_agent_action'):
                    callback.on_agent_action(action, **kwargs)

    # Initialize callback manager with our custom logger
    callback_manager = ExtendedCallbackManager([CrewAICallbackHandler()])
    ##############################################################################

    ##############################################################################
    class CustomCrew(Crew):
        def kickoff(self, inputs=None):
            """Override the kickoff method to add logging when agents perform tasks."""
            # print("\n🚀 Crew Execution Started...\n")
            log(f"🚀 Crew Execution Started...")

            # Ensure inputs is a dictionary (CrewAI requires it)
            inputs = inputs or {}

            # Run tasks manually, logging each agent's activity
            results = []
            for task in self.tasks:
                # print(f"\n🟢 Running Task: {task.description}")
                log(f"🟢 Running Task: {task.description}")
                # print(f"👤 Agent: {task.agent.role}")
                log(f"👤 Agent: {task.agent.role}")

                # Execute the task with agent and input
                result = task.execute()  # Removed the 'inputs' argument here

                # print(f"✅ Task Completed: {task.description}")
                # log(f"✅ Task Completed: {task.description}")
                log(f"✅ Task Completed")
                # print(f"📄 Output: {result}\n")
                log(f"📄 Output: {result}")

                results.append(result)
            # print("\n✅ Crew Execution Completed!\n")
            log(f"✅ Crew Execution Completed!")
            return results
    ##############################################################################

    ##############################################################################
    class LoggingAgent(Agent):
        """Custom Agent that introduces a delay during execution."""

        def execute(self, task):
            # callback_manager.on_agent_action(action=f"{self.role} is executing task: {task.description}")
            # log(f"🤖 {self.role} is executing task: {task.description}")
            log(f"🤖 {self.role} is executing the task")

            # Simulate final task completion
            time.sleep(2)  # **Delay before showing task completion**
            result = super().execute_task(task)  # Run the task

            # log(f"✅ {self.role} completed task. Output: {result}") # Already reported in the Crew section

            return result
    ##############################################################################

    ##############################################################################
    class LoggingTask(Task):
        """Custom Task that introduces a delay after execution."""

        def execute(self):  # Added the execute method
            """Execute the task using the assigned agent with proper logging."""

            # Use the actual task description as the serialized task name
            serialized_task_name = self.description

            # Use self.inputs if available, else provide an empty dictionary
            task_inputs = self.inputs if hasattr(self, "inputs") and self.inputs else {}

            # Log task start
            # callback_manager.on_chain_start(serialized=serialized_task_name, inputs=task_inputs)
            # log(f"📌 Task Started: {self.description}")
            log(f"📌 Task Started")

            # **Introduce a delay before executing**
            time.sleep(2)  # 2-second delay before execution starts

            # Call the agent's execute method to perform the task
            result = self.agent.execute(self)

            # **Introduce a delay before logging completion**
            time.sleep(3)  # 3-second delay before task completion logs

            # log(f"🎯 Task Completed: {self.description} | Output: {result}") # Already reported in the Crew section
            # callback_manager.on_chain_end(outputs=result)

            return result  # Return the result
    ##############################################################################

    ##############################################################################
    def generate_docx(result):
        doc = Document()
        doc.add_heading('Modeling Pipeline Automation - Fraud Undersampling', 0)
        doc.add_paragraph(result)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
    ##############################################################################

    ##############################################################################
    def get_download_link(bio, filename):
        b64 = base64.b64encode(bio.read()).decode()
        return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Modeling Documentation</a>'
    ##############################################################################

    # st.set_page_config(
    #     layout="wide"
    # )


    # Title
    # st.title("Modeling Pipeline Automation - Fraud prediction use case..")
    st.header("Modeling Pipeline Automation - Fraud prediction use case..")


    # Text Inputs
    dir_path = st.text_input("Enter project directory", "e.g., fraud/")
    # dir_path = st.text_input("Enter project directory", "fraud_und/")
    csv_name = st.text_input("Enter dataset name", "e.g., fraud.csv")
    train_data = st.text_input("Enter train-set name", "e.g., train.csv")
    test_data = st.text_input("Enter test-set name", "e.g., test.csv")
    # test_data = st.text_input("Enter test-set name", "test.csv")
    params = st.text_input("Enter filename for the hyperparameters", "e.g., hyper_params.txt")
    target = st.selectbox('Select target variable name', ('Class', 'TARGET', 'loan_status'))
    exclude_features = st.text_input("Enter the features to exclude", "e.g., Time, Date")


    directory_path= dir_path
    model_name= "model.joblib"
    train_data2= "train2.csv"
    test_data2= "test2.csv"

    manager = LoggingAgent(
        role="Data Science Manager",
        goal=(
            "Manage the agents effectively with these specific instructions:\n"
            "  Exploratory data analysis: Conduct a detailed exploratory data analysis on the provided dataset.\n"
            "  Feature Engineering: 1) Label encode categorical features, 2) Fix the issue of missing values, 3) Perform Random undersampling if class imbalance exists. \n"
            "  Model Selection: Select the best machine learning model based on on EDA insights and data distribution.\n"
            "  Hyperparameter tuning: Conduct hyperparameter tuning of the selected model using the 40% data provided.\n"
            "  Model Training: Train the selected model using the best hyperparameters derived earlier.\n"
            "  Model Evaluation: Evaluated the selected modela and include all metrics including top capture rate.\n"
        ),

        backstory="You're an experienced data science manager, skilled in overseeing complex fraud modeling projects and guiding teams to success. Your role is to coordinate the efforts of the crew members, ensuring that each task is completed on time and to the highest standard.",
        allow_delegation=True,
        max_iter=20,
        llm = llm,
        verbose=True,
        # memory=True,
        # tools=[eda_tool, mod_sel_tool, rag_tool,code_execution_tool], # Manager agent should not have tools
        cache=True
    )

    # Custom function to dynamically set and pass instructions
    def manager_task_logic(task):
        return {
            "eda_instruction": {"description": (
                "The Exploratory Data Analysis (EDA) report must include two things.\n"
                "1. Whether there are any categorical features present in the dataset.\n"
                "2. Whether the class/target distribution in the dataset is balanced or imbalanced.\n"
            ),
            "expected_output": (
            "The Final output should be a detailed EDA report with Title, Sub titles."
            )},
            "feature_engineer_instruction": {"description": (
                "Label encode categorical features if they exist.\n"
                "Fix missing value issues if they exists.\n"
                "Apply random undersampling to downsample the data if class imbalance is detected.\n"
            ),
            "expected_output": (
            "The Final output MUST be a report with Title and subtitles showing all the types of transformation performed on the dataset and the location of the transformed data."
            )},
            "model_selection_instruction": {"description": (
                "The model selection should only focus on:\n"
                "       1. Logistic regression\n"
                "       2. Decision tree\n"
                "       3. Support Vector Machine\n"
                "       4. Naive Bayes\n"
                "       5. Random forest\n"
                "       6. XGBoost\n"
                "       7. CatBoost - set 'verbose=0'\n"
                "       8. CatBoost\n"
                "       9. AdaBoosting\n\n"
            ),
            "expected_output": (
            "The Final output MUST be a report with Title and subtitles showing the selected machine learning model, and the rationale behind choosing the model."
            )},
            "hyperparameter_tuning_instruction":{"description": (
                "Use Grid search method ONLY on the selected model for the hyperparameter tuning. Make sure to use only numeric features for this exercise.\n"
            ),
            "expected_output": (
            "The Final output MUST be a report with Title and subtitles showing the best hyperparameters. Please print out the best hyperparameters."
            )},
            "model_training_instruction": {"description": (
              "Use the best hyperparameters from the 'Senior Machine Learning Engineer - Hyperparameter Tuning' to train the selected model.\n "
            ),
            "expected_output": (
            "The Final output MUST be a report with Title and subtitles showing the location of the saved model, train data and test data respectively."
            )},
            "evaluation_instruction":{"description": (
                "Calculate all evaluation metrics including accuracy, precision, recall, "
                "F1-score, and AUC.\n"
            ),
            "expected_output": (
            "The Final output MUST be a report with Title and subtitles showing values of all the evaluation metrics like accuracy, F1-score, recall, precision, and auc."
            )}
        }[task]


    eda_agent = LoggingAgent(
        role="Senior Data Scientist I",
        goal=(
            f"Conduct a detailed Exploratory data analysis on the provided dataset located at '{directory_path}{csv_name}'. "
            "The Final Answer should be a detailed EDA report with title and subtitles "
            ),
        backstory=(
            "You are a 'Senior Data Scientist - EDA' with expertise in conducting exploratory data analysis."
            ),
        verbose=True,
        allow_delegation=False,
        max_iter=5,
        tools=[eda_tool],
        # tools=[code_execution_tool],
        llm = llm,
        memory=True,
        cache=True
    )


    feature_engineer = LoggingAgent(
        role="Senior Data Scientist II",
        goal=(
            f"Preprocess the test and train dataset located on '{directory_path}{test_data}' and '{directory_path}{train_data}' respectively.\n"
            "To be effective in performing feature engineering, "
            "kindly follow the procedure below and handle the test and train data seperately:\n"
            "Feature Engineering on both test and train dataset:\n"
            f"  1) Drop the features found in '{exclude_features}' on both the train and test dataset\n"
            # "  2) Perform label encoding only on the categorical features with fit_transform and transform methods respectively. \n"
            # "  3) Use K-Nearest Neighbors (KNN) imputation to fill missing values on the train and test dataset with fit_transform and transform methods respectively. \n"
            f"  2) Check the unique values in the target variable named {target} for the test and train dataset separatley to determine if the proportion of any of the classes is equal to or greater than 60%'. "
            " if their is class imbalance on any(test or train dataset), perform random undersampling(random state = 42) on the test and train dataset seperately using fit_resample method. \n"
            f"  3) Save the new transformed test and train data in the directory '{directory_path}' as '{test_data2}' and '{train_data2}' respectively.\n"
            "Additionally, include the following in the generated code:\n"
            "  1. Create a variable 'result' with the first 3 rows of the transformed train data.\n"
            "The Final Answer MUST be a report with Title and subtitles showing all the types of transformation performed on the dataset and the location of the transformed data."
            ),
        backstory=(
            "You are a 'Senior Data Scientist - Feature Engineering' with specialization in Feature Engineering. You are responsible for understanding an EDA report and using the EDA findings for data preprocessing with the code_execution_tool."
            ),
        verbose=True,
        allow_delegation=False,
        max_iter=5,
        tools=[code_execution_tool],
        llm = llm,
        memory=True,
    )


    model_selection_agent = LoggingAgent(
        role="Machine Learning Engineer I",
        goal=(
            f"Select the best machine learning model using the GridSearchCV method on the dataset located at '{directory_path}{train_data2}'. "
            f"The target variable is named '{target}'.\n"
            "Give the rationale for the selected model."
            ),
        backstory=(
            "You are a 'Machine Learning Engineer - Model Selection' with expertise in selecting the best machine learning model based using GridSearchCV."
            ),
        verbose=True,
        allow_delegation=False,
        max_iter=5,
        tools=[mod_sel_tool],
        llm = llm,
        memory=True,
        cache=True
    )

    hyperparameter_tuning_agent = LoggingAgent(
        role="Senior Machine Learning Engineer I",
        goal=(
            "Tune only the hyperparameters of the selected "
            "model from the model selection agent. "
            f"Use only 40% of the dataset located at '{directory_path}{train_data2}' for this exercise. "
            "To be effective in choosing the parameters and values, "
            "kindly use the parameter for tuning the selected model below:\n"
            "    If Logistic Regression was selected, use 'C': [0.01, 0.1, 1, 10, 100], 'solver': ['lbfgs', 'liblinear']\n"
            "    If Decision Trees was selected, use  'max_depth': [1, 5, 10, 20], 'min_samples_split': [2, 5, 10]\n"
            "    If Support Vector Machine was selected, use  'C': [0.1, 1, 10], 'gamma': ['scale', 'auto']\n"
            "    If Random Forest was selected, use  'n_estimators': [50, 100, 200, 300], 'max_depth': [None, 1, 5, 10, 20]\n"
            "    If XGBoost was selected, use  'n_estimators': [100, 200], 'learning_rate': [0.01, 0.1], 'max_depth': [3, 5]\n"
            "    If CatBoost was selected, use 'verbose':0, 'iterations': [100, 200], 'learning_rate': [0.01, 0.1], 'depth': [3, 5]\n"
            "    If Gradient Boosting was selected, use  'n_estimators': [100, 200], 'learning_rate': [0.01, 0.1], 'max_depth': [3, 5]\n"
            "    If Ada Boost was selected, use  'n_estimators': [50, 100], 'learning_rate': [0.01, 0.1]\n"
            f"The target feature is named '{target}'. "
            "The random state and/or seed should be set to 42. \n"
            "The variable name for the best parameters must be called 'result' \n"
            f"Save the result as a txt file with the name '{params}' in the directory '{directory_path}'."
            ),
        backstory=(
            "You are a 'Machine Learning Engineer - Hyperparameter Tuning' with expertise in conducting hyperparameter tuning."
            ),
        verbose=True,
        allow_delegation=False,
        max_iter=5,
        tools=[code_execution_tool],
        llm = llm,
        memory=True,
        cache=True
    )

    train_model_agent = LoggingAgent(
        role="Senior Machine Learning Engineer II",
        goal=(
            "Train the selected model from the model selection agent.\n\n"
            # "Train the selected model from the model selection agent and NOT Random Forest.\n\n"
            "DATA:\n"
            f"* The train and test datasets are located at '{directory_path}{train_data2}' and '{directory_path}{test_data2}' respectively.\n"
            f"* The target variable is named '{target}'.\n"
            "HYPERPARAMETERS:\n"
            "* Please make sure to use the best hyperparameters to train the selected model.\n"
            "TRAINING:\n"
            "* Train the selected machine learning model using the train dataset and the best hyperparameters.\n"
            "SAVING:\n"
            f"* Save the trained model in the directory '{directory_path}' as '{model_name}'.\n"
            "* Create a dictionary variable called 'result' that contains the location of the saved model.\n\n"
            "IMPORTING joblib:\n"
            "* To import joblib, use 'import joblib'.\n\n"
            "ADDITIONAL INSTRUCTIONS:\n"
            "* Include any additional instructions from the manager to be able to complete the task.\n\n"
            "The Final Answer MUST be a report with Title and subtitles showing the location of the saved model, train data and test data respectively."

            ),
        backstory=(
            "You are a 'Senior Machine Learning Engineer - Model Training' with expertise in writing and executing Python codes to "
            "train the selected model using the best hyperparameter."
            ),
        verbose=True,
        allow_delegation=False,
        max_iter=5,
        tools=[code_execution_tool],
        llm = llm,
        memory=True,
        cache=True
    )

    evaluation_agent = LoggingAgent(
        role="Senior Machine Learning Engineer III",
        goal=(
            f"Evaluate the performance of the model located at '{directory_path}{model_name}' "
            f"The test data for this evaluation is located at '{directory_path}{test_data2}'. "
            f"The target variable is '{target}'. "
            "Within the generated code, a dict variable called 'result' should be created. "
            "The result variable should contain all the performance metrics accuracy, F1-score, recall, precision, and auc. "
            "Print the result variable"
            "The Final Answer MUST be a report with Title and subtitles showing values of all the evaluation metrics like accuracy, F1-score, recall, precision, and auc."
            ),
        backstory=(
            "You are a 'Senior Machine Learning Engineer - Model Evaluation' with expertise in writing and executing Python code for evaluating the "
            "a machine learning model. The evaluation includes accuracy, F1-score, precision, recall, AUC and overall performance."
            ),
        verbose=True,
        allow_delegation=False,
        max_iter=5,
        tools=[code_execution_tool],
        llm = llm,
        memory=True,
        cache=True
    )


    eda_task = LoggingTask(
        description=(
            f"Conduct a detailed Exploratory data analysis on the provided dataset located at '{directory_path}{csv_name}'."
            # "The report should show the shape of the dataset, missing values, the descriptive statistics "
            # "of the numeric features, correlation across features, feature skewness and data distribution.\n"
            # "Do not generating visual representations such as correlation matrices or histograms displaying data distributions, "
            # "the essential focus lies in providing a concise list of correlated features along with their respective skewness of "
            # "each feature e.g (positively skewed or negatively skewed). \n"
            # "The final output containing data shape, descriptive statistics, feature correlations, feature skewness , "
            # "and feature data types should be stored in a variable name called 'result'. "
            ),
        expected_output=(
            "A detailed EDA report with title and subtitles showing the shape of the dataset, "
            "missing values, the descriptive statistics of the numeric features, correlation "
            "across features, feature skewness and data distribution."
            ),
        # tools=[eda_tool],
        tools=[code_execution_tool],
        context=[manager_task_logic('eda_instruction')],
        agent=eda_agent
    )


    feature_task = LoggingTask(
        description=(
            f"Preprocess the test and train dataset located on '{directory_path}{test_data}' and '{directory_path}{train_data}' respectively.\n"
            # "To be effective in performing feature engineering, "
            # "kindly follow the procedure below and handle the test and train data seperately:\n"
            # "Feature Engineering on both test and train dataset:\n"
            # f"  1) Drop the features found in '{exclude_features}' on both the train and test dataset\n"
            # # "  2) Perform label encoding only on the categorical features with fit_transform and transform methods respectively. \n"
            # # "  3) Use K-Nearest Neighbors (KNN) imputation to fill missing values on the train and test dataset with fit_transform and transform methods respectively. \n"
            # f"  2) Check the unique values in the target variable named {target} for the test and train dataset separatley to determine if the proportion of any of the classes is equal to or greater than 60%'. "
            # " if their is class imbalance on any(test or train dataset), perform random undersampling(random state = 42) on the test and train dataset seperately using fit_resample method. \n"
            # f"  3) Save the new transformed test and train data in the directory '{directory_path}' as '{test_data2}' and '{train_data2}' respectively.\n"
            "Additionally, include the following in the generated code:\n"
            "  1. Create a variable 'result' with the first 3 rows of the transformed train data.\n"
            ),
        expected_output=(
            "A detailed Feature Engineering report showing type of transformation performed on the dataset and the first 3 rows of the transformed dataframe. "
            ),
        context=[manager_task_logic('feature_engineer_instruction')],
        tools=[code_execution_tool],
        agent=feature_engineer
    )


    model_selection_task = LoggingTask(
        description=(
            f"Select the best machine learning model using the GridSearchCV method on the dataset located at '{directory_path}{train_data2}'. "
            # f"The target variable is named '{target}'.\n"
            # "Give the rationale for the selected model."
                    ),
        expected_output=(
            "A detailed Model Selection report with title, subtitles showing "
            "the best machine learning model for the dataset. "
            "The report should show the model name like this - 'model': 'Linear regression', and an "
            "detailed explanation why the model was selected."
            ),
        # tools=[code_execution_tool],
        tools=[mod_sel_tool],
        context=[manager_task_logic('model_selection_instruction')],
        agent=model_selection_agent,
    )


    hyperparameter_tuning_task = LoggingTask(
        description=(
            "Tune only the hyperparameters of the selected "
            "model from the model selection agent. \n"
            # f"Use only 40% of the dataset located at '{directory_path}{train_data2}' for this exercise. "
            # "To be effective in choosing the parameters and values, "
            # "kindly use the parameter for tuning the selected model below:\n"
            "    If Logistic Regression was selected, use 'C': [0.01, 0.1, 1, 10, 100], 'solver': ['lbfgs', 'liblinear']\n"
            # "    If Decision Trees was selected, use  'max_depth': [1, 5, 10, 20], 'min_samples_split': [2, 5, 10]\n"
            # "    If Support Vector Machine was selected, use  'C': [0.1, 1, 10], 'gamma': ['scale', 'auto']\n"
            # "    If Random Forest was selected, use  'n_estimators': [50, 100, 200, 300], 'max_depth': [None, 1, 5, 10, 20]\n"
            # "    If XGBoost was selected, use  'n_estimators': [100, 200], 'learning_rate': [0.01, 0.1], 'max_depth': [3, 5]\n"
            # "    If CatBoost was selected, use 'verbose':0, 'iterations': [100, 200], 'learning_rate': [0.01, 0.1], 'depth': [3, 5]\n"
            # "    If Gradient Boosting was selected, use  'n_estimators': [100, 200], 'learning_rate': [0.01, 0.1], 'max_depth': [3, 5]\n"
            # "    If Ada Boost was selected, use  'n_estimators': [50, 100], 'learning_rate': [0.01, 0.1]\n"
            # f"The target feature is named '{target}'. "
            # "The random state and/or seed should be set to 42. \n"
            # "The variable name for the best parameters must be called 'result'. \n"
            f"Save the result as a txt file with the name '{params}' in the directory '{directory_path}'."
            ),
        expected_output=(
            "A detailed Hyperparameter Tuning report with title, subtitles showing "
            "the best hyperparameters " #like - max_depth: 5, n_estimators: 100, and their importance."
            ),
        tools=[code_execution_tool],
        context = [model_selection_task],
        agent=hyperparameter_tuning_agent
    )


    train_model_task = LoggingTask(
        description=(
            "Train the selected model from the model selection agent.\n\n"
            "DATA:\n"
            f"* The train and test datasets are located at '{directory_path}{train_data2}' and '{directory_path}{test_data2}' respectively.\n\n"
            # f"* The target variable is named '{target}'.\n"
            # "HYPERPARAMETERS:\n"
            # "* Please make sure to use the best hyperparameters to train the selected model.\n"
            # "TRAINING:\n"
            # "* Train the selected machine learning model using the train dataset and the best hyperparameters.\n"
            # "SAVING:\n"
            # f"* Save the trained model in the directory '{directory_path}' as '{model_name}'.\n"
            # "* Create a dictionary variable called 'result' that contains the location of the saved model.\n\n"
            # # "IMPORTING joblib:\n"
            # # "* To import joblib, use 'import joblib'.\n\n"
            "ADDITIONAL INSTRUCTIONS:\n"
            "* Include any additional instructions from the manager to be able to complete the task.\n\n"
            ),
        expected_output=(
            "A detailed Model Training report with title, subtitles showing "
            "if the training concluded successfully and the location of the saved model, train data and test data respectively."
            ),
        tools=[code_execution_tool],
        context = [hyperparameter_tuning_task],
        agent=train_model_agent
    )


    evaluation_task = LoggingTask(
        description=(
            f"Evaluate the performance of the model located at '{directory_path}{model_name}' "
            f"The test data for this evaluation is located at '{directory_path}{test_data2}'. \n"
            # f"The target variable is '{target}'. "
            # "Within the generated code, a dict variable called 'result' should be created. "
            # "The result variable should contain all the performance metrics accuracy, F1-score, recall, precision, top capture rate and auc. "
            # "Print the result variable"
            "Also include an additional instruction from the manager to complete the task. "
            ),
        expected_output=(
            "A detailed Model Evaluation report with title, subtitles showing "
            "all evaluation metrics like accuracy, F1-score, recall, precision, and auc."
            ),
        tools=[code_execution_tool],
        context=[manager_task_logic('evaluation_instruction')],
        agent=evaluation_agent
    )


    # Instantiate the crew with manager
    modeling_crew = CustomCrew(
        agents=[eda_agent,feature_engineer,model_selection_agent,hyperparameter_tuning_agent, train_model_agent, evaluation_agent],
        tasks=[eda_task,feature_task,model_selection_task,hyperparameter_tuning_task, train_model_task, evaluation_task],
        manager_llm=llm,
        manager_agent=manager,
        process=Process.hierarchical,
        verbose=True,
        memory=True,
        cache=True
    )



    modeling_inputs = {
      "csv_name": csv_name,
      "directory_path": dir_path,
      "model_name": "model.joblib",
      "train_data": train_data,
      "test_data": test_data,
      "train_data2": "train2.csv",
      "test_data2": "test2.csv",
      "params": params,
      "target": target,
      "exclude_features": exclude_features,
    }


    # Execution
    if st.button("Run the Modeling workflow"):
        with st.spinner('Performing machine learning modeling...'):
            result = modeling_crew.kickoff(inputs=modeling_inputs)
            st.write(result)
            docx_file = generate_docx(result)

            download_link = get_download_link(docx_file, "fraud_modeling_documentation.docx")

            st.markdown(download_link, unsafe_allow_html=True)

